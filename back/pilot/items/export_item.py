from html import escape
from io import BytesIO

from docx import Document as DocxDocument

from django.utils.translation import ugettext_lazy as _
from django.http import HttpResponse, HttpResponseServerError
from django.shortcuts import get_object_or_404
from django.template.loader import render_to_string

from pilot.items.models import EditSession, Item
from pilot.utils import export_utils
from pilot.utils.alpha import to_alpha
from pilot.utils.prosemirror.html_docx import add_html as html_docx_add_html
from pilot.utils.pdf import render_to_pdf, PdfRenderingException
from pilot.utils.serialization import SerializationFormat
from settings.base import backend_path

EXPORT_PLAIN_TEXT = 'plaintext'
EXPORT_HTML = 'html'
EXPORT_PDF = 'pdf'
EXPORT_DOCX = 'docx'


class ItemContentExporter(object):
    def __init__(self, item, session_pk=None):
        if session_pk:
            self.session = get_object_or_404(
                EditSession.objects.select_related('created_by', 'restored_from'),
                pk=session_pk,
                item=item
            )
        else:
            self.session = item.last_session

    def export_to_text(self):
        fields = []
        for converter in self.session.converters.values():
            fields.append("{label}\n----\n{value}".format(
                label=converter.label,
                value=converter.serialize(SerializationFormat.TEXT))
            )
        return '\n========================\n'.join(fields)

    def export_to_html(self):
        return render_to_string(
            "items/export/exporthtml.html",
            {'item_content': self.session},
        )

    def export_to_docx(self):
        document = DocxDocument(backend_path('pilot', 'utils', 'prosemirror', 'html_docx', 'template.docx'))
        document.add_paragraph("Item ID = {}".format(self.session.item_id))
        document.add_paragraph("Version = {}".format(self.session.version))
        document.add_paragraph("------------------------------")

        for converter in self.session.converters.values():
            label = document.add_paragraph()
            html_docx_add_html(label, "<strong>{}</strong>".format(converter.label))

            html_content = converter.serialize(SerializationFormat.HTML)
            if html_content:
                content = document.add_paragraph()
                html_docx_add_html(content, html_content)

            document.add_paragraph()

        f = BytesIO()
        document.save(f)
        f.seek(0)
        return f.getvalue()

    def export_to_pdf(self):
        return render_to_pdf(
            'items/export/exportpdf.html',
            {
                'pagesize': 'A4',
                'item_content': self.session,
            }
        )

    def export_to_response(self, export_type):
        if export_type == EXPORT_PLAIN_TEXT:
            response_content = self.export_to_text()
            content_type = 'text/plain'
            file_extension = 'txt'

        elif export_type == EXPORT_HTML:
            response_content = self.export_to_html()
            content_type = 'text/html'
            file_extension = 'html'

        elif export_type == EXPORT_DOCX:
            response_content = self.export_to_docx()
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            file_extension = 'docx'

        elif export_type == EXPORT_PDF:
            try:
                response_content = self.export_to_pdf()
                content_type = 'application/pdf'
                file_extension = 'pdf'
            except PdfRenderingException as e:
                return HttpResponseServerError('We had some errors<pre>%s</pre>' % escape(e.html))

        else:
            raise ValueError('Incorrect export type : {}'.format(export_type))

        response = HttpResponse(response_content, content_type=content_type)
        response['Content-Disposition'] = 'attachment; filename="{}_{}.{}"'.format(
            to_alpha(self.session.title),
            self.session.get_version_display(),
            file_extension
        )
        return response


class ItemXLSExporter(export_utils.BaseXLSExporter):
    model = Item
    metadata_fields = export_utils.ITEM_METADATA_FIELDS

    def __init__(self, queryset, output_file, with_content=False):
        super(ItemXLSExporter, self).__init__(queryset, output_file)
        self.with_content=with_content

    def get_header(self):
        # 1/ Item content or title
        if self.with_content:
            yield _('Contenu')
        else:
            yield _('Titre')

        # 2/ Metadata fields
        yield from super(ItemXLSExporter, self).get_header()

        # 3/ Publication
        yield _('Date de publication prévue')
        yield _('Publication effective')

    def get_row(self, i, item):
        # 1/ Item content or title
        if self.with_content:
            yield ItemContentExporter(item).export_to_text()
        else:
            yield item.title

        # 2/ Metadata fields
        yield from super(ItemXLSExporter, self).get_row(i, item)

        # 3/ Publication
        if item.publication_task:
            yield export_utils.format_date(item.publication_task.deadline)
            yield export_utils.format_date_time(item.publication_task.done_at)

    def add_style(self):
        # Wrap the content/title column, and shift other wrapped columns by 1
        self.wrapped_columns = [1] + [i+1 for i in self.wrapped_columns]

        super(ItemXLSExporter, self).add_style()

        # Increase the width of the content/title column
        self.worksheet.column_dimensions['A'].width = 100 if self.with_content else 50
        # Decrease the width of the id column
        self.worksheet.column_dimensions['B'].width = 10
